import sys
import pandas as pd
from docx import Document
import math
from itertools import takewhile

doc = Document(sys.argv[1])
rows = []
current_year = None
current_section = None
current_content = []
CHUNK_SZ = 10_000  # python xlsx cutsoff at 32k chars silently
EXCLUDED_SECTIONS = ["Plan", "Weekly Summaries, starting on Monday inclusive of Sunday."]

for p in doc.paragraphs:
    style, text = p.style.name, p.text
    if not text.strip():
        continue
    if style == "Heading 2":
        current_year = text.strip()
    elif style == "Heading 4":
        if current_content and current_section not in EXCLUDED_SECTIONS:
            full = "\n".join(current_content)
            total_len = len(full)
            print(total_len, len(current_content), CHUNK_SZ)
            n_chunks = math.ceil(total_len / CHUNK_SZ)
            target = math.ceil(total_len / n_chunks)
            chunks = []
            buf = []
            buf_len = 0
            for line in current_content:
                line_len = len(line) + 1
                if buf and buf_len + line_len > target:
                    chunks.append(buf)
                    buf = []
                    buf_len = 0
                buf.append(line)
                buf_len += line_len
            if buf:
                chunks.append(buf)

            for i, chunk in enumerate(chunks, start=1):
                rows.append(
                    {
                        "Year": current_year,
                        "Section": f"{current_section} ({i})",
                        "Content": "\n".join(chunk),
                    }
                )
        current_section = text.strip()
        current_content = []
    else:
        if current_section and current_section not in EXCLUDED_SECTIONS:
            current_content.append(text.rstrip())

if current_section and current_section not in EXCLUDED_SECTIONS:
    rows.append(
        {"Year": current_year, "Section": current_section, "Content": "\n".join(current_content)}
    )

df = pd.DataFrame(rows, columns=["Year", "Section", "Content"])
df.to_excel(sys.argv[2], engine="openpyxl", index=False)


# %%
if True:
    print(1 / 0)

import sys
import csv
from docx import Document
import math

doc = Document(sys.argv[1])
rows = []
current_year = None
current_section = None
current_content = []

EXCLUDED_SECTIONS = ["Plan", "Weekly Summaries, starting on Monday inclusive of Sunday."]

for p in doc.paragraphs:
    s = p.style.name
    t = p.text.strip()
    if not t:
        continue

    if s == "Heading 2":
        current_year = t
    elif s == "Heading 4":
        # ignore first data at start before journal
        if current_section and current_content and current_section not in EXCLUDED_SECTIONS:
            rows.append(
                {
                    "Year": current_year,
                    "Section": current_section,
                    "Content": "\n".join(current_content),
                }
            )
        current_section = t
        current_content = []
    else:
        # Only add content if we have all required fields and section is not excluded
        if current_section not in EXCLUDED_SECTIONS:
            current_content.append(t)

# Add the last section if it exists
if current_section and current_content and current_section not in EXCLUDED_SECTIONS:
    rows.append(
        {
            "Year": current_year,
            "Section": current_section,
            "Content": "\n".join(current_content),
        }
    )

# Write to CSV with proper quoting
with open(sys.argv[2], "w", newline="", encoding="utf-8") as f:
    w = csv.DictWriter(
        f, fieldnames=["Year", "Section", "Content"], quoting=csv.QUOTE_ALL, escapechar="\\"
    )
    w.writeheader()
    w.writerows(rows)
